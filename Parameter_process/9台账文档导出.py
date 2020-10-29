# -*- coding: utf-8 -*-#
# -------------------------------------------------------------------------------
# Author:       A
# Date:         2020-09-25
# -------------------------------------------------------------------------------
import os
import time
import re
import xlrd
from docx import Document
from docx.shared import Pt
from docx.enum.section import WD_ORIENTATION, WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx import shared
from docx.shared import RGBColor


# 写入地块情况说明
def doc_(*args):
    # 添加文本
    document.styles['Normal'].font.name = 'Times New Roman'
    document.styles['Normal'].element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋')
    document.add_paragraph()
    document.add_paragraph()
    document.add_paragraph()
    document.add_paragraph()
    bianhao = message[0].value
    qx = message[1].value
    xz = message[2].value
    xzc = message[3].value
    lon = message[4].value
    lat = message[5].value
    zllb = message[6].value
    rwmj = message[26].value
    sd19area = message[28].value
    sd20area = message[30].value
    ym19area = message[32].value
    ym20area = message[34].value
    sd19pz = message[27].value
    sd20pz = message[29].value
    ym19pz = message[31].value
    ym20pz = message[33].value
    qt19 = message[36].value
    qt20 = message[37].value

    pztz = message[7].value
    jght = message[11].value
    sfg = message[12].value
    yhsf = message[14].value
    tghl = message[18].value
    bgfgd = message[19].value
    sgmg = message[20].value
    zzjgtz = message[21].value
    title_P = document.add_paragraph()
    title_P.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_T = title_P.add_run(f'{bianhao}图斑情况说明')
    fontT = title_T.font
    fontT.size = Pt(18)
    fontT.name = '仿宋'
    document.add_paragraph()

    # title_T.element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋')
    # 2019年作物种植情况
    def qingkuang2019(ym19area, sd19area, qt19):
        qingkuang19nian = ""
        ym19nian = ""
        if ym19area == 0:
            ym19nian = ""
        else:
            ym19nian = f"玉米品种：{ym19pz}，共{ym19area}亩,"
        ym19nian.strip()
        sd19nian = ""
        if sd19area == 0:
            sd19nian = ""
        else:
            sd19nian = f"水稻品种：{sd19pz}，共{sd19area}亩,"
        sd19nian.strip()
        qt19nian = ""
        if qt19.replace(' ', '') == '':
            qt19nian = "无其它作物。"
        else:
            qt19nian = f"其它主要作物：{qt19},"
        qt19nian.strip()
        if ym19area == 0 and sd19area == 0 and qt19.replace(' ', '') == '':
            qingkuang19nian = "该地块2019年种植情况为:无农作物。"
        else:
            qingkuang19nian = "该地块2019年种植情况为:" + sd19nian + ym19nian + qt19nian
        return qingkuang19nian

    qingkuang2019(ym19area, sd19area, qt19)

    # 2020年作物种植情况
    def qingkuang2020(ym20area, sd20area, qt20):
        qingkuang20nian = ""
        ym20nian = ""
        if ym20area == 0:
            ym20nian = ""
        else:
            ym20nian = f"玉米品种：{ym20pz}，共{ym20area}亩,"
        ym20nian.strip()
        sd20nian = ""
        if sd20area == 0:
            sd20nian = ""
        else:
            sd20nian = f"玉米品种：{sd20pz}，共{sd20area}亩,"
        sd20nian.strip()
        qt20nian = ""
        if qt20.replace(' ', '') == '':
            qt20nian = "无其它作物。"
        else:
            qt20nian = f"其它主要作物：{qt20},"
        qt20nian.strip()
        if ym20area == 0 and sd20area == 0 and qt20.replace(' ', '') == '':
            qingkuang20nian = "该地块2020年种植情况为:无农作物。"
        else:
            qingkuang20nian = "该地块2020年种植情况为:" + sd20nian + ym20nian + qt20nian
        return qingkuang20nian

    paragraph_neirong = document.add_paragraph()
    run5 = paragraph_neirong.add_run(
        f"    编号{bianhao}地块位于{qx}{xz}{xzc},地块中心坐标：东经{lon},北纬{lat},质量类别为{zllb},地块总面积{rwmj}亩,经现场调查后，%s%s" % (
            qingkuang2019(ym19area, sd19area, qt19), qingkuang2020(ym20area, sd20area, qt20)))
    paragraph_qingkuang = document.add_paragraph()
    # 措施落实情况
    cqcs_qk = ''
    cqcs_dict = {'品种调整': pztz, '秸秆还田': jght, '深翻耕': sfg, '优化施肥': yhsf, '退耕还林还草': tghl, '土地利用变更为非农用地': bgfgd,
                 '农户自主休耕': sgmg, '种植结构调整': zzjgtz}
    ix = [x + ':' + cqcs_dict[x] + '亩' for x in cqcs_dict if cqcs_dict[x] != ""]

    for one_mes in ix:
        one_mes = str(one_mes).replace(' ', '') + '，'
        cqcs_qk += one_mes
    cqcs_qk = cqcs_qk[:-1] + '。'
    run1 = paragraph_qingkuang.add_run(
        f"    图斑措施落实情况为：{cqcs_qk}")

    # 措施佐证文件情况
    ix_1 = [x for x in cqcs_dict if cqcs_dict[x] != ""]
    zzwj_dict = {'品种调整': '佐证资料文件见附件第一章P1-35', '深翻耕': '佐证资料文件见附件第四章P126-222',
                 '优化施肥': '佐证资料文件见附件第四章P126-222', '退耕还林还草': '佐证资料文件见附件第三章P77-116',
                 '土地利用变更为非农用地': '佐证资料文件见附件第四章P75', '种植结构调整': '佐证文件资料见附件第一章P37-59',
                 '秸秆还田': '佐证文件资料见', '农户自主休耕': '佐证文件见附件第五章P225-234'}
    sm_list = [x + "措施" + zzwj_dict[x] + "," for x in ix_1]
    zzwj_sm = ''
    for one_cs in sm_list:
        zzwj_sm += one_cs
    zzwj_sm = zzwj_sm[:-1] + '。'
    paragraph_wenjian = document.add_paragraph()
    run6 = paragraph_wenjian.add_run(f"    {zzwj_sm}")
    paragraph_shuoming = document.add_paragraph()
    run2 = paragraph_shuoming.add_run("特此说明：")
    document.add_paragraph()
    paragraph_shiti = document.add_paragraph()
    run3 = paragraph_shiti.add_run("%s" % ((" ") * 46))
    document.add_paragraph()
    paragraph_shijian = document.add_paragraph()
    run4 = paragraph_shijian.add_run("%s责任单位（盖章）：" % ((" ") * 36))

    # 设置字体格式及行距
    for i in [run1, run2, run3, run4, run5, run6]:
        i.line_spacing = 1.5  # 设置行距
        font = i.font  # 建立对象
        font.size = Pt(14)  # 设置字号
        font.name = '仿宋'  # 设置字体
        i.element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋')  # 设置字体
        # run.blod = True  # 字体加粗
    # 增加分页
    document.add_page_break()
    print(f"写入{message[0].value}")
    return document


# 写入平面布置图
def pingmianbuzhitu(*args):
    # section = document.add_section(start_type=WD_SECTION_START.CONTINUOUS)
    # section.orientation = WD_ORIENTATION.LANDSCAPE  # 设置横向（为插入图片做准备）
    # page_h, page_w = section.page_width, section.page_height
    # section.page_width = page_w
    # section.page_height = page_h
    document.add_picture("%s" % tupia_dict.get(message[0].value),
                         width=shared.Cm(17.2), height=shared.Cm(22.8))
    # document.add_page_break()
    # section = document.add_section(start_type=WD_SECTION_START.CONTINUOUS)
    # section.orientation = WD_ORIENTATION.PORTRAIT  # 设置纵向（还原，为插入文档做准备）
    # page_h, page_w = section.page_width, section.page_height
    # section.page_width = page_w
    # section.page_height = page_h


# 写入调查表
def diaochabiao(**args):
    title = document.add_heading('', 3)
    title1 = title.add_run(f'{message[0].value}')
    title1.font.color.rgb = RGBColor(0, 0, 0)
    document.add_picture("%s" % diaochabiao_dict.get(message[0].value),
                         width=shared.Cm(17), height=shared.Cm(22))  # 写入调查表


# 写入现场图片
def xianchang_pictures(*args):
    try:
        for pictures_xc in os.listdir(xctp_dict.get(*args)):
            if pictures_xc.endswith('jpg'):
                in_pictures_xc = os.path.join(xctp_dict.get(*args), pictures_xc)
                document.add_picture("%s" % in_pictures_xc,
                                     height=shared.Cm(11), width=shared.Cm(16))  # 写入现场图片
    except AttributeError:
        print("无该地块照片")


def dikuai_shuoming(*args):
    document.add_picture("%s" % shuoming_dict.get(message[0].value),
                         width=shared.Cm(15), height=shared.Cm(21))  # 写入地块说明


sheet_path = r"F:\1台账导出文档基础资料\湄潭县台账\4湄潭县表格信息\MTX20200908_dissolve.xls"  # 表格位置
pingmiantu_path = r"F:\1台账导出文档基础资料\湄潭县台账\3湄潭县平面布置图_旋转"  # 平面图位置

xianchangtu_path = r"F:\1台账导出文档基础资料\湄潭县台账\5湄潭县现场图片"  # 现场图位置

diaochabiao_path = r"F:\1台账导出文档基础资料\湄潭县台账\1湄潭县调查表cut_image"  # 调查表位置

# dikuaishuoming_path = r"F:\湄潭县台账\2湄潭县情况说明表"  # 情况说明表位置

document_save_path = r'F:\2台账导出结果文档\湄潭县'

pdf_dir = []
diaochabiao_dict = {}
for roots, dirs, files in os.walk(diaochabiao_path):  # 收集全部调查表图片到字典
    for one_file in files:
        if str(one_file).endswith("jpg"):
            dk_key = re.findall(r"5\d{11}", str(one_file))[0]
            diaochabiao_dict[dk_key] = os.path.join(roots, one_file)
tupia_dict = {}
for roots, dirs, files in os.walk(pingmiantu_path):  # 收集全部平面布置图到字典
    for one_file in files:
        if str(one_file).endswith("jpg"):
            dk_key = re.findall(r"5\d{11}", str(one_file))[0]
            tupia_dict[dk_key] = os.path.join(roots, one_file)
xctp_dict = {}  # 获取地块编码对应现场图片路径
for xctp_dir in os.listdir(xianchangtu_path):
    print("")
    xctp_dict[xctp_dir] = os.path.join(xianchangtu_path, xctp_dir)
# shuoming_dict = {}
# for roots, dirs, files in os.walk(dikuaishuoming_path):
#     for one_file in files:
#         if str(one_file).endswith("jpg"):
#             dk_key = re.findall(r"5\d{11}", str(one_file))[0]
#             shuoming_dict[dk_key] = os.path.join(roots, one_file)
work_book = xlrd.open_workbook(sheet_path)  # 打开工作薄
work_sheet = work_book.sheet_by_index(1)  # 打开表格

xz_list = []
message_list = []  # 全县总地块信息列表
for one_row in range(1, work_sheet.nrows):  # 逐一写入各个地块情况及资料
    message1 = work_sheet.row(one_row)  # 每个地块的信息
    message_list.append(message1)
    xz_list.append(message1[2].value)
    xz_list = list(set(xz_list))  # 全县乡镇列表
print(xz_list)
time.sleep(5)
for one_xz in xz_list:
    print(one_xz)
    try:
        os.makedirs(os.path.join(r'%s' % document_save_path, "%s" % one_xz))
    except:
        print("文件夹已创建")
    one_message_list = list(filter(lambda x: x[2].value == one_xz, message_list))
    one_message_list.sort(key=lambda x: x[2].value)
    cz_list = []
    for one_row in one_message_list:  # 逐一写入各个地块情况及资料
        cz_list.append(one_row[3].value)
        cz_list = list(set(cz_list))
    print(cz_list)
    for one_cz in cz_list:
        if os.path.exists(r"%s\%s\%s.docx" % (document_save_path, one_xz, one_cz)):
            print("文档已存在")
        else:
            document = Document()  # 创建空文档
            title1 = document.add_heading('', 2)
            title2 = title1.add_run(f'{one_cz}')
            title2.font.color.rgb = RGBColor(0, 0, 0)
            print(one_cz)
            two_message_list = list(filter(lambda x: x[3].value == one_cz, one_message_list))
            two_message_list.sort(key=lambda x: x[3].value)
            for message in two_message_list:
                print(message)
                print(message[0].value)
                diaochabiao()  # 写入调查表
                doc_(message)  # 写入地块说明
                pingmianbuzhitu()  # 写入平面布置图
                xianchang_pictures(message[0].value)  # 写入现场图片
            document.save(r'%s\%s\%s.docx' % (document_save_path, one_xz, one_cz))
            del document
            print(one_cz)
            time.sleep(3)
