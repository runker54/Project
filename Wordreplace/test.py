from openpyxl import load_workbook
from docx import Document
import time
import os

# def replace_text_with_format(doc, old_text, new_text):
#     for paragraph in doc.paragraphs:
#         if old_text in paragraph.text:
#             # 获取旧文字所在段落的格式
#             run = None
#             for run in paragraph.runs:
#                 if old_text in run.text:
#                     break
#             font = run.font if run else None
#             style = paragraph.style

#             # 替换文字
#             inline = paragraph.runs
#             for i in range(len(inline)):
#                 if old_text in inline[i].text:
#                     inline[i].text = inline[i].text.replace(old_text, new_text)

#             # 应用旧文字的格式到新文字
#             for run in paragraph.runs:
#                 if new_text in run.text:
#                     run.font.bold = font.bold if font else False
#                     run.font.italic = font.italic if font else False
#                     run.font.underline = font.underline if font else False
#                     run.font.color.rgb = font.color.rgb if font else None
#                     paragraph.style = style
#                     break
#     # 遍历文档的所有表格
#     for table in doc.tables:
#         for row in table.rows:
#             for cell in row.cells:
#                 if old_text in cell.text:
#                     # 获取旧文字所在单元格的格式
#                     paragraphs = cell.paragraphs
#                     font = None
#                     style = None

#                     for paragraph in paragraphs:
#                         if old_text in paragraph.text:
#                             runs = paragraph.runs
#                             for run in runs:
#                                 if old_text in run.text:
#                                     font = run.font
#                                     style = paragraph.style
#                                     break
#                             if font and style:
#                                 break
                    
#                     # 替换文字
#                     for paragraph in paragraphs:
#                         if old_text in paragraph.text:
#                             for run in paragraph.runs:
#                                 if old_text in run.text:
#                                     run.text = run.text.replace(old_text, new_text)
#                                     run.font.bold = font.bold if font else False
#                                     run.font.italic = font.italic if font else False
#                                     run.font.underline = font.underline if font else False
#                                     run.font.color.rgb = font.color.rgb if font else None
#                                     break
                    
#                     # 应用旧文字的格式到新文字
#                     for paragraph in paragraphs:
#                         if new_text in paragraph.text:
#                             paragraph.style = style
#                             break
def replace_text_with_format(doc, old_text, new_text):
    for paragraph in doc.paragraphs:
        if old_text in paragraph.text:
            for run in paragraph.runs:
                if old_text in run.text:
                    # 获取旧文字所在段落的格式
                    font = run.font if run else None
                    style = paragraph.style

                    # 替换文字并应用格式
                    inline = paragraph.runs
                    for i in range(len(inline)):
                        if old_text in inline[i].text:
                            inline[i].text = inline[i].text.replace(old_text, new_text)
                            inline[i].font.bold = font.bold if font else False
                            inline[i].font.italic = font.italic if font else False
                            inline[i].font.underline = font.underline if font else False
                            inline[i].font.color.rgb = font.color.rgb if font else None
                    paragraph.style = style

    # 遍历文档的所有表格
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if old_text in cell.text:
                    # 获取旧文字所在单元格的格式
                    paragraphs = cell.paragraphs
                    font = None
                    style = None

                    for paragraph in paragraphs:
                        for run in paragraph.runs:
                            if old_text in run.text:
                                font = run.font
                                style = paragraph.style
                                break
                        if font and style:
                            break

                    # 替换文字并应用格式
                    for paragraph in paragraphs:
                        for run in paragraph.runs:
                            if old_text in run.text:
                                run.text = run.text.replace(old_text, new_text)
                                run.font.bold = font.bold if font else False
                                run.font.italic = font.italic if font else False
                                run.font.underline = font.underline if font else False
                                run.font.color.rgb = font.color.rgb if font else None
                    for paragraph in paragraphs:
                        if new_text in paragraph.text:
                            paragraph.style = style
                            break


if __name__ == '__main__':
    excel_file = 'test.xlsx'  # Excel文件路径
    word_template = 'test.docx'  # Word模板文件路径
    save_path = 'savepath'  # 保存路径
    # checkmark_hex_code = '\u221A'  # √符号
    # dui ='\u221A'
    dui = '\u221A' # 对钩
    # cuo = '\u2611' # 方框
    cuo = '\u25A1' # 方框
    new_text = dui
    doc = Document(word_template)
    # 保存修改后的文档
    wb = load_workbook(excel_file)  # 打开excel文件
    ws = wb.active  # 获取当前活跃的表单
    for row in ws.iter_rows():

        xian = str(row[3].value)
        xiang = str(row[4].value)
        cun = str(row[5].value)
        diwz = str(row[6].value)
        tbbh = str(row[7].value)
        zxzb = str(row[8].value)
        hasnone = str(row[9].value)
        area = str(row[10].value)
        doc = Document(word_template)
        # message_list = [xian, xiang, diwz, tbbh, zxzb, area]
        # target_list = ["XXXX","YYYY","WZ","BH","ZB","MJ"]
        replace_text_with_format(doc, "XXXX", xian)
        replace_text_with_format(doc, "YYYY", xiang)
        replace_text_with_format(doc, "HB", diwz)
        replace_text_with_format(doc, "HS", tbbh)
        replace_text_with_format(doc, "HC", zxzb)
        replace_text_with_format(doc, "HD", area)
        if hasnone == '是':
            replace_text_with_format(doc, 'HA', dui)
            replace_text_with_format(doc, 'NO', cuo)
        else:
            replace_text_with_format(doc, 'HA', cuo)
            replace_text_with_format(doc, 'NO', dui)
        doc.save(f"./savepath/{diwz}{tbbh}.docx")
        print("done")
        