#!/usr/bin/env python
# -*- coding: utf-8 -*-
# Date:2020-12-10
# Author:Runker54
# -----------------------
from tqdm import tqdm
import win32com.client as win32
from docx import Document
from docx.shared import Pt
from docx.enum.section import WD_ORIENTATION, WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx import shared
from docx.shared import RGBColor
import time
import re
import os
import shutil

quxian = '紫云苗族布依族自治县'
temp = r'F:\%s\temp' % quxian  # 过程数据
merge = r'F:\%s\merge' % quxian   # 保存位置

for one_z in os.listdir(merge):
    if '$' not in one_z:
        z_path = os.path.join(merge, one_z)
        print(one_z)
        doc = Document(z_path)
        np = doc.paragraphs[0].insert_paragraph_before()
        p = np.add_run(one_z[:one_z.rfind('.')])
        font = p.font  # 建立对象
        font.size = Pt(12)  # 设置字号
        font.name = '仿宋'  # 设置字体
        p.font.color.rgb = RGBColor(250, 20, 0)
        p.element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋')  # 设置字体
        doc.add_page_break()
        doc.save(z_path)
    else:
        pass