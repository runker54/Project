#!/usr/bin/env python
# -*- coding: utf-8 -*-
# Date:2020-12-10
# Author:Runker54
# -----------------------
from tqdm import tqdm
import win32com.client as win32
from docx import Document
from docx.shared import Pt
from docx.enum.section import WD_ORIENTATION, WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx import shared
from docx.shared import RGBColor
import time
import re
import os
import shutil
adress = r'E:\20201210\核算\铜仁市'
for x_name in os.listdir(adress):
    quxian = x_name
    root_path = os.path.join(adress, x_name)  # 源数据
    root_adress = f'F:'
    qxname = root_path[root_path.rfind('\\') + 1:]
    array_path = os.path.join(root_adress, qxname)
    try:
        os.makedirs(os.path.join(root_adress, qxname))
        os.makedirs(os.path.join(root_adress, 'temp'))
        os.makedirs(os.path.join(root_adress, 'merge'))
    except:
        pass
    temp = os.path.join(root_adress, 'temp')
    merge = os.path.join(root_adress, 'merge')

    for one_xz in os.listdir(root_path):
        if os.path.isfile(r'%s\%s.docx' % (merge, one_xz)):
            pass
        else:
            xz_path = os.path.join(root_path, one_xz)
            print(one_xz)
            docx_list = [os.path.join(xz_path, x) for x in os.listdir(xz_path)]
            temp_path = os.path.join(temp, one_xz)
            temp1_path = os.path.join(merge, one_xz)
            try:
                os.makedirs(temp_path)
            except IOError:
                pass
            for file_paht in docx_list:
                print(file_paht)
                doc = Document(file_paht)
                # 核算点位匹配规则 r'5\d{5}.{3,}\d{3,}'
                # 普通编码规则 r'5\d{9,}'
                name = re.findall(r'5\d{5}.{3,}\d{3,}', file_paht)[0]
                np = doc.paragraphs[0].insert_paragraph_before()
                p = np.add_run(name)
                font = p.font  # 建立对象
                font.size = Pt(11)  # 设置字号
                font.name = '仿宋'  # 设置字体
                p.font.color.rgb = RGBColor(250, 0, 0)
                p.element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋')  # 设置字体
                doc.add_page_break()
                doc.save(os.path.join(temp_path, name + '.docx'))
            # 打开word软件
            word = win32.gencache.EnsureDispatch('Word.Application')
            # 非可视化运行
            word.Visible = False
            output = word.Documents.Add()  # 新建合并后空白文档
            files1 = []
            for roots, dirs, files in os.walk(temp_path):
                for file in files:
                    if file.endswith('docx'):
                        files1.append(os.path.join(roots, file))
                    else:
                        pass
            for file in files1:
                if len(files1) == 1:
                    shutil.copyfile(file, os.path.join(merge, '%s.docx' % one_xz))
                    continue
                else:
                    output.Application.Selection.Range.InsertFile(file)  # 拼接文档
            output.SaveAs(r'%s\%s.docx' % (merge, one_xz))  # 保存
            output.Close()
            word.Quit()
    shutil.move(temp, array_path)
    shutil.move(merge, array_path)
