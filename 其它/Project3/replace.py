#!/usr/bin/env python
# -*- coding:utf-8 -*-
# Time:2020/4/9
from docx import Document
import xlrd
import xlwt


def chang_text(old_text, new_text):
    all_paragraphs = document.paragraphs
    for paragraph in all_paragraphs:
        for run in paragraph.runs:
            run_text = run.text.replace(old_text, new_text)
            run.text = run_text
    all_tables = document.tables
    for table in all_tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        cell_text = cell.text.replace(old_text, new_text)
                        run.text = cell_text


wb = xlrd.open_workbook("信息.xls")

document = Document("安全利用类模板.docx")

sheet = wb.sheet_by_index(0)
for table_row in range(1, sheet.nrows):
    for table_col in range(0, sheet.ncols):
        chang_text(str(sheet.cell_value(0, table_col)), str(sheet.cell_value(table_row, table_col)))

    document.save("%s%s统计信息.docx" % (str(sheet.cell_value(table_row, 0)), str(sheet.cell_value(table_row, 5))))
