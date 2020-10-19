# -*- coding: utf-8 -*-#
# -------------------------------------------------------------------------------
# Author:       A
# Date:         2020-09-26
# -------------------------------------------------------------------------------
import os
import time
import re
import xlrd
from docx import Document
from docx.shared import Pt
from docx.enum.section import WD_ORIENTATION, WD_SECTION_START
from docx.oxml.ns import qn
from docx import shared
from docx.shared import RGBColor
from PIL import Image
import fitz


def get_file(*args):
    for roots, dirs, files in os.walk(*args):
        for file in files:
            if os.path.splitext(file)[1] == '.pdf':  # 目录下包含.pdf的文件
                pdf_dir.append(os.path.join(roots, file))
    print(len(pdf_dir))


# 调查表PDF转图片
def conver_img(*args):
    for pdf in pdf_dir:
        doc = fitz.open(pdf)
        pdf_name = os.path.splitext(pdf)[0]
        for pg in range(doc.pageCount):
            page = doc[pg]
            rotate = int(0)
            # 每个尺寸的缩放系数为2，这将为我们生成分辨率提高四倍的图像。
            zoom_x = 2.0
            zoom_y = 2.0
            trans = fitz.Matrix(zoom_x, zoom_y).preRotate(rotate)
            pm = page.getPixmap(matrix=trans, alpha=False)
            if os.path.exists(str(pdf).replace("pdf", "jpg")):
                print("图片已存在")
            else:
                pm.writePNG('%s.jpg' % pdf_name)
                print(pdf)


# 写入地块情况说明
def doc_(*args):
    # 加入不同等级的标题
    # title = document.add_heading(f'{message[0].value}', 1)
    # title = document.add_heading('', 1)
    # title1 = title.add_run(f'{message[0].value}')
    # title1.font.color.rgb = RGBColor(250, 250, 250)
    # 添加文本
    document.add_paragraph()
    document.add_paragraph()
    bianhao = message[0].value
    qx = message[1].value
    xz = message[2].value
    xzc = message[3].value
    lon = message[4].value
    lat = message[5].value
    zllb = message[6].value
    rwmj = message[26].value
    sd19area = message[28].value
    sd20area = message[30].value
    ym19area = message[32].value
    ym20area = message[34].value
    sd19pz = message[27].value
    sd20pz = message[29].value
    ym19pz = message[31].value
    ym20pz = message[33].value
    qt19 = message[36].value
    qt20 = message[37].value

    pztz = message[7].value
    jght = message[11].value
    sfg = message[12].value
    yhsf = message[14].value
    tghl = message[18].value
    bgfgd = message[19].value
    sgmg = message[20].value
    zzjgtz = message[21].value
    paragraph_neirong = document.add_paragraph()
    run5 = paragraph_neirong.add_run(
        f"  编号{bianhao}地块位于{qx}{xz}{xzc},地块中心坐标：东经{lon},北纬{lat},质量类别为{zllb},总任务面积{rwmj},经现场调查后，"
        f"该地块2019年种植情况为，玉米品种：{ym19pz}共 {ym19area} 亩，水稻品种：{sd19pz}共 {sd19area} 亩，其它主要作物为：{qt19}，"
        f"2020年种植情况为,玉米品种：{ym20pz}共 {ym20area} 亩，水稻品种：{sd20pz}共 {sd20area} 亩，其它主要作物为：{qt20}。")
    paragraph_qingkuang = document.add_paragraph()
    document.add_paragraph()
    run1 = paragraph_qingkuang.add_run(
        f"  图斑措施落实情况为：品种调整：{pztz} 亩，秸秆还田： {jght} 亩，深翻耕： {sfg} 亩，优化施肥： {yhsf} 亩，退耕还林还草： {tghl} 亩，"
        f"土地利用变更为非农用地： {bgfgd} 亩，少耕休耕免耕： {sgmg} 亩，种植结构调整： {zzjgtz} 亩。实施时间为2017年至2020年。")

    paragraph_shuoming = document.add_paragraph()
    run2 = paragraph_shuoming.add_run("特此说明：")
    document.add_paragraph()
    paragraph_shiti = document.add_paragraph()
    run3 = paragraph_shiti.add_run(f"%s{xzc}人民政府" % ((" ") * 38))
    document.add_paragraph()
    paragraph_shijian = document.add_paragraph()
    run4 = paragraph_shijian.add_run("%s   年    月    日" % ((" ") * 36))

    # 设置行距
    run5.line_spacing = 1.5
    run1.line_spacing = 1.5
    run2.line_spacing = 1.5
    run3.line_spacing = 1.5
    run4.line_spacing = 1.5

    # 设置字体格式
    for i in [run1, run2, run3, run4, run5]:
        font = i.font
        font.size = Pt(14)
        font.name = '仿宋'
        i._element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋')
        # run.blod = True

    document.add_paragraph()
    document.add_paragraph()
    document.add_paragraph()
    # 增加分页
    document.add_page_break()
    print(f"写入{message[0].value}")
    return document


# 写入平面布置图
def pingmianbuzhitu(*args):
    # im = Image.open(tupia_dict.get(message[0].value))
    # out = im.transpose(Image.ROTATE_90)
    # out.save(tupia_dict.get(message[0].value))
    document.add_picture("%s" % tupia_dict.get(message[0].value), height=shared.Cm(21.5), width=shared.Cm(15.5))
    # section = document.add_section(start_type=WD_SECTION_START.CONTINUOUS)
    # section.orientation = WD_ORIENTATION.LANDSCAPE  # 设置横向（为插入图片做准备）
    # page_h, page_w = section.page_width, section.page_height
    # section.page_width = page_w
    # section.page_height = page_h
    # document.add_picture("%s" % tupia_dict.get(message[0].value),
    #                      width=shared.Cm(21.5), height=shared.Cm(15.5))
    # document.add_page_break()
    # section = document.add_section(start_type=WD_SECTION_START.CONTINUOUS)
    # section.orientation = WD_ORIENTATION.PORTRAIT  # 设置纵向（还原，为插入文档做准备）
    # page_h, page_w = section.page_width, section.page_height
    # section.page_width = page_w
    # section.page_height = page_h


# 写入调查表
def diaochabiao(**args):
    title = document.add_heading('', 2)
    title1 = title.add_run(f'{message[0].value}')
    # title1.font.color.rgb = RGBColor(250, 250, 250)
    document.add_picture("%s" % diaochabiao_dict.get(str(message[0].value).strip()),
                         width=shared.Cm(16), height=shared.Cm(21.2))  # 写入调查表


# 写入现场图片
def xianchang_pictures(*args):
    for pictures_xc in os.listdir(xctp_dict.get(*args)):
        in_pictures_xc = os.path.join(xctp_dict.get(*args), pictures_xc)
        document.add_picture("%s" % in_pictures_xc, height=shared.Cm(11))  # 写入调查表


sheet_path = r"G:\金沙县台账\document文档\JSX_20200925_dissolve.xls"  # 表格位置
pingmiantu_path = r"G:\金沙县台账\document文档\金沙县图片"  # 平面图位置
# xianchangtu_path = r"C:\Users\65680\Desktop\复兴镇\复兴镇"  # 现场图位置
diaochabiao_path = r"G:\金沙县台账\document文档\金沙县JPG裁剪"  # 调查表位置

pdf_dir = []
# get_file(diaochabiao_path)  # 获取调查表PDF
# conver_img()  # 执行PDF转图片
diaochabiao_dict = {}
for roots, dirs, files in os.walk(diaochabiao_path):  # 收集全部调查表图片到字典
    for one_file in files:
        if str(one_file).endswith("jpg"):
            dk_key = re.findall(r"5\d{11}", str(one_file))[0]
            diaochabiao_dict[dk_key] = os.path.join(roots, one_file)
tupia_dict = {}
for roots, dirs, files in os.walk(pingmiantu_path):  # 收集全部平面布置图到字典
    for one_file in files:
        if str(one_file).endswith("jpg"):
            dk_key = re.findall(r"5\d{11}", str(one_file))[0]
            tupia_dict[dk_key] = os.path.join(roots, one_file)


xctp_dict = {}  # 获取地块编码对应现场图片路径
# for xctp_dir in os.listdir(xianchangtu_path):
#     xctp_dict[xctp_dir] = os.path.join(xianchangtu_path, xctp_dir)
work_book = xlrd.open_workbook(sheet_path)  # 打开工作薄
work_sheet = work_book.sheet_by_index(1)  # 打开表格

xz_list = []
message_list = []
for one_row in range(1, work_sheet.nrows):  # 逐一写入各个地块情况及资料
    message1 = work_sheet.row(one_row)  # 每个地块的信息
    message_list.append(message1)
    xz_list.append(message1[2].value)
    xz_list = list(set(xz_list))

print(xz_list)
time.sleep(5)
for one_xz in ['桂花乡', '平坝镇']:
    print(one_xz)
    try:
        os.makedirs(os.path.join(r'F:\JSX', "%s" % one_xz))
    except:
        print("文件夹已创建")

    one_message_list = list(filter(lambda x: x[2].value == one_xz, message_list))
    one_message_list.sort(key=lambda x: x[2].value)
    cz_list = []
    for one_row in one_message_list:  # 逐一写入各个地块情况及资料
        cz_list.append(one_row[3].value)
        cz_list = list(set(cz_list))
    print(cz_list)
    for one_cz in cz_list:
        if os.path.exists(r"F:\JSX\%s\%s.docx" % (one_xz, one_cz)):
            print("文档已存在")
        else:
            document = Document()  # 创建空文档
            title1 = document.add_heading('', 1)
            title2 = title1.add_run(f'{one_cz}')
            # title2.font.color.rgb = RGBColor(250, 250, 250)
            print(one_cz)
            two_message_list = list(filter(lambda x: x[3].value == one_cz, one_message_list))
            two_message_list.sort(key=lambda x: x[3].value)
            for message in two_message_list:
                print(message)
                print(diaochabiao_dict.get(str(message[0].value).strip()))
                diaochabiao()  # 写入调查表
                # doc_(message)  # 写入地块说明
                # xianchang_pictures(message[0].value)  # 写入现场图片
                pingmianbuzhitu()  # 写入平面布置图
            document.save(r'F:\JSX\%s\%s.docx' % (one_xz, one_cz))
            print(one_cz)
            time.sleep(3)
