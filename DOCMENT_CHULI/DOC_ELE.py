# coding: utf-8
import docx
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
doc = docx.Document()
paragraph = doc.add_paragraph()
paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
run = paragraph.add_run("")
run.add_picture(r'C:\Users\65680\Desktop\新建文件夹\520425000686.jpg')
doc.save(r"C:\Users\65680\Desktop\DOC\ccc.docx")
