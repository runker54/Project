# -*- coding: utf-8 -*-#
# -------------------------------------------------------------------------------
# Author:       A
# Date:         2020-10-14
# -------------------------------------------------------------------------------
import os
import time
import re
import xlrd
from docx import Document
from docx.shared import Pt
from docx.enum.section import WD_ORIENTATION, WD_SECTION_START
from docx.oxml.ns import qn
from docx import shared
from docx.shared import RGBColor
import fitz


def doc_(*args):
    # 加入不同等级的标题
    # title = document.add_heading(f'{message[0].value}', 1)
    # title = document.add_heading('', 1)
    # title1 = title.add_run(f'{message[0].value}')
    # title1.font.color.rgb = RGBColor(250, 250, 250)
    # 添加文本
    document.add_paragraph()
    document.add_paragraph()
    document.add_paragraph()
    document.add_paragraph()
    document.add_paragraph()
    bianhao = message[0].value
    qx = message[1].value
    xz = message[2].value
    xzc = message[3].value
    lon = message[4].value
    lat = message[5].value
    zllb = message[6].value
    rwmj = message[26].value
    sd19area = message[28].value
    sd20area = message[30].value
    ym19area = message[32].value
    ym20area = message[34].value
    sd19pz = message[27].value
    sd20pz = message[29].value
    ym19pz = message[31].value
    ym20pz = message[33].value
    qt19 = message[36].value
    qt20 = message[37].value

    pztz = message[7].value
    jght = message[11].value
    sfg = message[12].value
    yhsf = message[14].value
    tghl = message[18].value
    bgfgd = message[19].value
    sgmg = message[20].value
    zzjgtz = message[21].value

    # 2019年作物种植情况



    def qingkuang2019(ym19area, sd19area, qt19):
        qingkuang19nian = ""
        ym19nian = ""
        if ym19area == 0:
            ym19nian = ""
        else:
            ym19nian = f"玉米品种：{ym19pz}共 {ym19area} 亩,"
        ym19nian.strip()
        sd19nian = ""
        if sd19area == 0:
            sd19nian = ""
        else:
            sd19nian = f"水稻品种：{sd19pz}共 {sd19area} 亩,"
        sd19nian.strip()
        qt19nian = ""
        if qt19 == ' ':
            qt19nian = "无其它作物，"
        else:
            qt19nian = f"其它主要作物为：{qt19},"
        qt19nian.strip()
        if ym19area == 0 and sd19area == 0 and qt19 == ' ':
            qingkuang19nian = "该地块2019年种植情况为:无农作物"
        else:
            qingkuang19nian = "该地块2019年种植情况为:"+sd19nian + ym19nian + qt19nian
        return qingkuang19nian

    qingkuang2019(ym19area, sd19area, qt19)

    # 2020年作物种植情况
    def qingkuang2020(ym19area, sd19area, qt19):
        qingkuang20nian = ""
        ym20nian = ""
        if ym20area == 0:
            ym20nian = ""
        else:
            ym20nian = f"玉米品种：{ym20pz}共 {ym20area} 亩,"
        ym20nian.strip()
        sd20nian = ""
        if sd20area == 0:
            sd20nian = ""
        else:
            sd20nian = f"玉米品种：{sd20pz}共 {sd20area} 亩,"
        sd20nian.strip()
        qt20nian = ""
        if qt20 == ' ':
            qt20nian = "无其它作物"
        else:
            qt20nian = f"其它主要作物为：{qt20},"
        qt20nian.strip()
        if ym20area == 0 and sd20area == 0 and qt20 == ' ':
            qingkuang20nian = "该地块2020年种植情况为:无农作物"
        else:
            qingkuang20nian = "该地块2020年种植情况为:"+sd20nian + ym20nian + qt20nian
        return qingkuang20nian

    paragraph_neirong = document.add_paragraph()
    run5 = paragraph_neirong.add_run(
        f"  编号{bianhao}地块位于{qx}{xz}{xzc},地块中心坐标：东经{lon},北纬{lat},质量类别为{zllb},总任务面积{rwmj}亩,经现场调查后，%s,%s" % (
        qingkuang2019(ym19area, sd19area, qt19), qingkuang2020(ym20area, sd20area, qt20)))
    paragraph_qingkuang = document.add_paragraph()
    document.add_paragraph()
    run1 = paragraph_qingkuang.add_run(
        f"  图斑措施落实情况为：品种调整：{pztz} 亩，秸秆还田： {jght} 亩，深翻耕： {sfg} 亩，优化施肥： {yhsf} 亩，退耕还林还草： {tghl} 亩，"
        f"土地利用变更为非农用地： {bgfgd} 亩，少耕休耕免耕： {sgmg} 亩，种植结构调整： {zzjgtz} 亩。实施时间为2017年至2020年。")

    paragraph_shuoming = document.add_paragraph()
    run2 = paragraph_shuoming.add_run("特此说明：")
    document.add_paragraph()
    paragraph_shiti = document.add_paragraph()
    run3 = paragraph_shiti.add_run("%s" % ((" ") * 46))
    document.add_paragraph()
    paragraph_shijian = document.add_paragraph()
    run4 = paragraph_shijian.add_run("%s   年    月    日" % ((" ") * 36))

    # 设置行距
    run5.line_spacing = 1.5
    run1.line_spacing = 1.5
    run2.line_spacing = 1.5
    run3.line_spacing = 1.5
    run4.line_spacing = 1.5

    # 设置字体格式
    for i in [run1, run2, run3, run4, run5]:
        font = i.font
        font.size = Pt(14)
        font.name = '仿宋'
        i._element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋')
        # run.blod = True

    document.add_paragraph()
    document.add_paragraph()
    document.add_paragraph()
    # 增加分页
    document.add_page_break()
    print(f"写入{message[0].value}")
    return document


sheet_path = r"C:\Users\65680\Desktop\ZYX20200909_dissolve.xls"  # 表格位置

work_book = xlrd.open_workbook(sheet_path)  # 打开工作薄
work_sheet = work_book.sheet_by_index(1)  # 打开表格

xz_list = []
message_list = []
for one_row in range(1, work_sheet.nrows):  # 逐一写入各个地块情况及资料
    message1 = work_sheet.row(one_row)  # 每个地块的信息
    message_list.append(message1)
    xz_list.append(message1[2].value)
    xz_list = list(set(xz_list))

print(xz_list)
time.sleep(5)
for one_xz in xz_list:
    print(one_xz)
    try:
        os.makedirs(os.path.join(r'C:\Users\65680\Desktop\地块说明测试', "%s" % one_xz))
    except:
        print("文件夹已创建")

    one_message_list = list(filter(lambda x: x[2].value == one_xz, message_list))
    one_message_list.sort(key=lambda x: x[2].value)
    cz_list = []
    for one_row in one_message_list:  # 逐一写入各个地块情况及资料
        cz_list.append(one_row[3].value)
        cz_list = list(set(cz_list))
    print(cz_list)
    for one_cz in cz_list:
        if os.path.exists(r"C:\Users\65680\Desktop\地块说明测试\%s\%s.docx" % (one_xz, one_cz)):
            print("文档已存在")
        else:
            document = Document()  # 创建空文档
            title1 = document.add_heading('', 2)
            title2 = title1.add_run(f'{one_cz}')
            # title2.font.color.rgb = RGBColor(250, 250, 250)
            print(one_cz)
            two_message_list = list(filter(lambda x: x[3].value == one_cz, one_message_list))
            two_message_list.sort(key=lambda x: x[3].value)
            for message in two_message_list:
                # print(message)
                doc_(message)
            document.save(r'C:\Users\65680\Desktop\地块说明测试\%s\%s.docx' % (one_xz, one_cz))
            print(one_cz)
            time.sleep(3)
