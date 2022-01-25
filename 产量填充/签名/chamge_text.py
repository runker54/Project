#!/usr/bin/env python
# -*- coding: utf-8 -*-
# Date:2021-11-03
# Author:Runker54
# ----------------------

name_path = r"C:\Users\65680\Desktop\name_path"
doc_path = r"C:\Users\65680\Desktop\生成Word文档\目录.docx"
from docx import Document
from docx.shared import Inches,Cm,Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import datetime
import re
import os
import time
docx = Document(doc_path)
tables = docx.tables
tables_number = len(tables)
for one_table in range(0, tables_number, 2):
    table = tables[one_table]
    print(table)
    # 地块编号
    name1_cahe = table.cell(2, 3).text
    # 图片1
    pic_1 = table.cell(22, 0).text
    print(pic_1)
    # name1 = name1_cahe[1].strip()
    run1 = table.cell(21, 0).paragraphs[0].add_run("")
    image1 = run1.add_picture(r"C:\Users\65680\Desktop\111.jpg")
    image1.width = Cm(18)
    image1.height = Cm(15)
    # # 采样成员
    # name2_cahe = str(table.cell(0, 1).text).split("：")
    # name2 = name2_cahe[1].strip()
    # run2 = table.cell(1, 1).paragraphs[0].add_run("")
    # image2 = run2.add_picture(os.path.join(name_path, name2+".png"))
    # image2.width = Cm(1.8)
    # image2.height = Cm(0.5)

#     print(name1)
docx.save(r"C:\Users\65680\Desktop\skdfl.docx")