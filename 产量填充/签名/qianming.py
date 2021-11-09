#!/usr/bin/env python
# -*- coding: utf-8 -*-
# Date:2021-11-03
# Author:Runker54
# ----------------------

name_path = r"C:\Users\65680\Desktop\name_path"
doc_path = r"C:\Users\65680\Desktop\玉屏县2021年受污染耕地安全利用项目采样记录表.docx"
from docx import Document
from docx.shared import Inches,Cm,Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import datetime
import re
import os
import time
docx = Document(doc_path)
tables = docx.tables
tables_number = len(tables)
for one_table in range(2, tables_number, 2):
    table = tables[one_table]
    # 负责人
    name1_cahe = str(table.cell(0, 0).text).split("：")
    name1 = name1_cahe[1].strip()
    run1 = table.cell(1, 0).paragraphs[0].add_run("")
    image1 = run1.add_picture(os.path.join(name_path, name1+".png"))
    image1.width = Cm(1.8)
    image1.height = Cm(0.5)
    # 采样成员
    name2_cahe = str(table.cell(0, 1).text).split("：")
    name2 = name2_cahe[1].strip()
    run2 = table.cell(1, 1).paragraphs[0].add_run("")
    image2 = run2.add_picture(os.path.join(name_path, name2+".png"))
    image2.width = Cm(1.8)
    image2.height = Cm(0.5)

    print(name1)
docx.save(r"C:\Users\65680\Desktop\已签名\skdfl.c")